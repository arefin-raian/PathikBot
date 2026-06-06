"""PDF pre-processor: temp copy → -1pt fonts → page-spacing adjustment → PDF conversion with Aspose."""
import os
import shutil
import zipfile
import hashlib
import time
from pathlib import Path
from lxml import etree
from docx import Document as DocxDocument
from docx.shared import Pt

import jpype
import jpype.imports

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _unique_name() -> str:
    raw = f"{time.time()}{os.urandom(8).hex()}"
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


def _find_jvm_dll() -> str:
    java_home = os.environ.get("JAVA_HOME") or os.environ.get("JDK_HOME")
    if java_home:
        for p in (Path(java_home) / "jre" / "bin" / "server" / "jvm.dll",
                  Path(java_home) / "bin" / "server" / "jvm.dll"):
            if p.is_file():
                return str(p)
    for base in (Path("C:/Program Files/Eclipse Adoptium"),
                 Path("C:/Program Files/Java")):
        if base.is_dir():
            for jdk_dir in base.iterdir():
                c = jdk_dir / "bin" / "server" / "jvm.dll"
                if c.is_file():
                    return str(c)
    import subprocess
    try:
        out = subprocess.run(["where", "java"], capture_output=True, text=True, timeout=5).stdout
        if out:
            home = Path(out.strip().splitlines()[0]).resolve().parent.parent
            c = home / "bin" / "server" / "jvm.dll"
            if c.is_file():
                return str(c)
    except Exception:
        pass
    raise RuntimeError("No jvm.dll found. Set JAVA_HOME to your JDK path.")


def _ensure_jvm(jar_path: str):
    if not jpype.isJVMStarted():
        jpype.startJVM(_find_jvm_dll(), classpath=[jar_path], convertStrings=True)


def _read_docx_xml(docx_path: str) -> bytes:
    with zipfile.ZipFile(docx_path, "r") as z:
        return z.read("word/document.xml")


def _write_docx_xml(docx_path: str, xml_bytes: bytes):
    """Replace word/document.xml inside the zip with new xml_bytes."""
    entries = {}
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            entries[name] = z.read(name)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            if name == "word/document.xml":
                zout.writestr(name, xml_bytes)
            else:
                zout.writestr(name, data)


def get_page_count(docx_path: str) -> int:
    jar = str(Path(__file__).resolve().parent.parent / "aspose-words-20.12-jdk17-cracked.jar")
    _ensure_jvm(jar)
    from com.aspose.words import Document as AsposeDocument
    doc = AsposeDocument(docx_path)
    return doc.getPageCount()


def copy_docx(src: str) -> str:
    stem = Path(src).stem
    name = f"{stem}_pdfadj_{_unique_name()}.docx"
    dst = str(Path(src).parent / name)
    shutil.copy2(src, dst)
    return dst


def reduce_font_sizes(docx_path: str):
    doc = DocxDocument(docx_path)
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.size is not None:
                run.font.size = run.font.size - Pt(1)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.size is not None:
                            run.font.size = run.font.size - Pt(1)
    doc.save(docx_path)


def adjust_page_spacing(docx_path: str):
    """Per page-boundary adjustment: try 2 spacers, then 1, then 0."""
    orig_pages = get_page_count(docx_path)
    if orig_pages <= 1:
        return

    xml_bytes = _read_docx_xml(docx_path)
    root = etree.fromstring(xml_bytes)
    body = root.find(f"{{{W}}}body")
    if body is None:
        return

    # Find positions between consecutive tables
    children = list(body)
    boundaries = []
    for i in range(len(children) - 1):
        if children[i].tag == f"{{{W}}}tbl" and children[i + 1].tag == f"{{{W}}}tbl":
            boundaries.append(i + 1)

    spacer_xml = (
        f'<w:p xmlns:w="{W}"><w:pPr>'
        f'<w:spacing w:line="480" w:lineRule="auto" w:before="120" w:after="120"/>'
        f'</w:pPr></w:p>'
    )
    spacer_elem = etree.fromstring(spacer_xml.encode())

    for idx in reversed(boundaries):
        # Try adding 2 spacers at this boundary
        for _ in range(2):
            body.insert(idx, etree.fromstring(etree.tostring(spacer_elem)))
        _write_docx_xml(docx_path, etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True))
        pages = get_page_count(docx_path)
        if pages <= orig_pages:
            continue

        # Too much — remove one spacer
        children2 = list(body)
        spacers_at_boundary = [c for c in children2 if c.tag == f"{{{W}}}p" and
                               not c.findall(f".//{{{W}}}t")]
        for sp in spacers_at_boundary[:1]:
            body.remove(sp)
        _write_docx_xml(docx_path, etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True))
        pages = get_page_count(docx_path)
        if pages <= orig_pages:
            continue

        # Still too much — remove the other spacer
        children3 = list(body)
        spacers_at_boundary = [c for c in children3 if c.tag == f"{{{W}}}p" and
                               not c.findall(f".//{{{W}}}t")]
        for sp in spacers_at_boundary[:1]:
            body.remove(sp)
        _write_docx_xml(docx_path, etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True))


def convert_docx_to_pdf(docx_path: str, pdf_path: str, fonts_dir: str):
    jar = str(Path(__file__).resolve().parent.parent / "aspose-words-20.12-jdk17-cracked.jar")
    _ensure_jvm(jar)
    from com.aspose.words import Document as AsposeDocument, SaveFormat, FontSettings

    font_settings = FontSettings()
    if os.path.isdir(fonts_dir):
        font_settings.setFontsFolder(fonts_dir, True)

    doc = AsposeDocument(docx_path)
    doc.setFontSettings(font_settings)
    doc.save(pdf_path, SaveFormat.PDF)


def prepare_for_pdf(original_docx: str) -> tuple:
    """Full pipeline: copy → reduce fonts → adjust spacing.
    Returns (temp_path, pdf_path).
    """
    temp = copy_docx(original_docx)
    reduce_font_sizes(temp)
    adjust_page_spacing(temp)
    pdf = str(Path(original_docx).with_suffix(".pdf"))
    return temp, pdf


def cleanup_temp(temp_path: str):
    try:
        if os.path.isfile(temp_path):
            os.remove(temp_path)
    except Exception:
        pass
