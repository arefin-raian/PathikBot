import os
from docx import Document
from datetime import datetime, date
import calendar
from dotenv import load_dotenv
from core.expense_calculations import calculate_summary
from docx_generator.bijoy_converter import convert_to_bijoy
from docx_generator.docx_xml_helpers import set_cell_text
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import copy

load_dotenv()

MONTHS_BN = {
    1: "Rvbyqvix", 2: "†deªæqvix", 3: "gvP©", 4: "GwcÖj",
    5: "†g", 6: "Ryb", 7: "RyjvB", 8: "AvM÷",
    9: "†m‡Þ¤^i", 10: "A‡±vei", 11: "b‡f¤^i", 12: "wW‡m¤^i"
}

class LogsheetGenerator:
    def __init__(self, template_path="templates/Logsheet Template.docx"):
        self.template_path = template_path
        # We'll load the template to get the table structures
        self.base_doc = Document(template_path)
        self.header_table_tpl = self.base_doc.tables[0]
        self.type1_table_tpl = self.base_doc.tables[1]
        self.type2_table_tpl = self.base_doc.tables[2]
        self.type3_table_tpl = self.base_doc.tables[3]
        self.summary_table_tpl = self.base_doc.tables[4]

    def _clone_table(self, table_tpl, parent_element):
        new_tbl = copy.deepcopy(table_tpl._tbl)
        parent_element.append(new_tbl)
        # Find the new table object in the document
        # This is a bit hacky but works for adding to the end
        return new_tbl

    def _add_page_break(self, doc):
        doc.add_page_break()

    def generate_report(self, entries, month, year, output_path):
        summary = calculate_summary(entries)
        new_doc = Document()
        
        # Set landscape orientation
        section = new_doc.sections[0]
        section.orientation = 1 # WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        # Set margins (0.5 inch = 720 DXA)
        section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = 457200 # 0.5 inch in EMU

        body = new_doc._body._body

        # 1. Type 1 Page
        # Add Header Table
        header_tbl = self._clone_table(self.header_table_tpl, body)
        from docx.table import Table
        header_table_obj = Table(header_tbl, new_doc)
        self._fill_header(header_table_obj, month, year)

        # Add Type 1 Data Table
        type1_tbl = self._clone_table(self.type1_table_tpl, body)
        type1_table_obj = Table(type1_tbl, new_doc)
        
        page1_entries = entries[:3]
        for i, entry in enumerate(page1_entries):
            self._fill_data_row(type1_table_obj.rows[i+1], entry, i)
        
        # 2. Type 2 Pages
        remaining_entries = entries[3:]
        
        # If we have more than can fit in Type 3 (let's say Type 3 takes at most 3 entries)
        # We need Type 2 pages for entries between index 3 and (total - 3 or something)
        # Actually, let's follow the 4-per-page rule for middle pages.
        
        # We need to save at least 1 entry for the last page (Type 3)
        # If total entries <= 3, Type 1 is enough (but prompt says Type 3 and 4 always exist)
        
        # Let's say: 
        # Page 1: 1-3
        # Middle Pages: 4-7, 8-11, ...
        # Last Data Page: whatever is left (1-4 entries) + Total row.
        
        # Determine how many entries go to Type 3
        # Usually 1 to 4 entries.
        if len(entries) > 3:
            num_middle_entries = len(entries) - 3
            # Reserve at least 1 and at most 4 for the last page
            last_page_count = num_middle_entries % 4
            if last_page_count == 0: last_page_count = 4
            
            middle_entries_count = num_middle_entries - last_page_count
            middle_entries = entries[3:3+middle_entries_count]
            last_entries = entries[3+middle_entries_count:]
            
            # Add Type 2 Pages
            for i in range(0, len(middle_entries), 4):
                self._add_page_break(new_doc)
                t2_tbl = self._clone_table(self.type2_table_tpl, body)
                t2_table_obj = Table(t2_tbl, new_doc)
                chunk = middle_entries[i:i+4]
                for j, entry in enumerate(chunk):
                    self._fill_data_row(t2_table_obj.rows[j+1], entry, 3 + i + j)
            
            # Add Type 3 Page
            self._add_page_break(new_doc)
            t3_tbl = self._clone_table(self.type3_table_tpl, body)
            t3_table_obj = Table(t3_tbl, new_doc)
            for j, entry in enumerate(last_entries):
                self._fill_data_row(t3_table_obj.rows[j+1], entry, 3 + middle_entries_count + j)
            # Fill Total Row (Last row of Table 3)
            self._fill_total_row(t3_table_obj.rows[-1], summary)
        else:
            # Special case: very few entries. Just add an empty Type 3.
            self._add_page_break(new_doc)
            t3_tbl = self._clone_table(self.type3_table_tpl, body)
            t3_table_obj = Table(t3_tbl, new_doc)
            self._fill_total_row(t3_table_obj.rows[-1], summary)

        # 3. Blank Type 2 Page
        self._add_page_break(new_doc)
        self._clone_table(self.type2_table_tpl, body)

        # 4. Type 4 Page (Summary)
        self._add_page_break(new_doc)
        summary_tbl = self._clone_table(self.summary_table_tpl, body)
        summary_table_obj = Table(summary_tbl, new_doc)
        self._fill_summary(summary_table_obj, summary)

        new_doc.save(output_path)
        return output_path

    def _fill_header(self, table, month, year):
        # Company Info from .env
        company = os.getenv('COMPANY_NAME', 'বিএআই এগিকালচারাল ইন্ডাস্ট্রিজ লিমিটেড')
        officer = os.getenv('OFFICER_NAME', 'মো: আশরাফ আলী')
        desig = os.getenv('DESIGNATION', 'টেরিটরি মার্কেটিং অফিসার')
        posting = os.getenv('POSTING_AREA', 'ডোমার')
        brand = os.getenv('MOTORCYCLE_BRAND', 'বাজাজ ডিসকভার')
        depo = os.getenv('DEPO_NAME', 'রংপুর')

        # Row 1: Company Name (Left), Month (Right)
        # Based on Table 0 inspection, we need to find exact cells.
        # This is template dependent. I'll use the most likely cells.
        month_name = MONTHS_BN.get(int(month), "")
        month_year_str = f"{month_name}/{year} Bs"
        
        # Header Table (Table 0) structure from prompt:
        # Row 1: [0,0] Company, [0, 1] Month
        # Row 2: [1,0] Depo, [1,1] Brand, [1,2] Date (Last day)
        # Row 3: [2,0] Name, [2,1] Desig, [2,2] Area
        
        # I will set the text using Bijoy conversion
        set_cell_text(table.cell(0, 0), f"†Kvñ·vbxi bvg: {company}") # "কোম্পানির নাম: "
        set_cell_text(table.cell(0, 1), f"gv‡mi bvg: {month_year_str}") # "মাসের নাম: "
        
        set_cell_text(table.cell(1, 0), f"wW‡cv bvg: {depo}")
        set_cell_text(table.cell(1, 1), f"‡gvUi mvB‡K‡ji bª¨vÛ: {brand}")
        
        last_day = calendar.monthrange(int(year), int(month))[1]
        last_date_str = f"{last_day:02d}/{int(month):02d}/{year} Bs"
        set_cell_text(table.cell(1, 2), f"ZvwiL: {last_date_str}")
        
        set_cell_text(table.cell(2, 0), f"Awdmv‡ii bvg: {officer}")
        set_cell_text(table.cell(2, 1), f"c`ex: {desig}")
        set_cell_text(table.cell(2, 2), f"†cvw÷s GjvKv: {posting}")

    def _fill_data_row(self, row, entry, index):
        # Index is 0-based
        set_cell_text(row.cells[0], f"{index+1:02d}")
        set_cell_text(row.cells[1], datetime.strptime(entry['date'], '%Y-%m-%d').strftime('%d/%m/%y'))
        
        if entry.get('entry_type') == 'MONTHLY_MEETING':
            # Special formatting for meeting
            venue = convert_to_bijoy(entry.get('venue', ''))
            fee = entry.get('transport_fee', 0)
            text = f"{venue}│\n†Wvgvi nB‡Z {venue} evm I A‡Uv‡Z hvZvqvZ fvov={fee}/-"
            set_cell_text(row.cells[2], text)
            set_cell_text(row.cells[3], f"{entry['odo_start']}")
            set_cell_text(row.cells[4], f"{entry['odo_end']}")
            set_cell_text(row.cells[5], "00")
            set_cell_text(row.cells[10], f"{fee}/-")
            set_cell_text(row.cells[11], "gvwmK wgwUs") # "মাসিক মিটিং"
        else:
            # Format names: strip brackets for DOCX
            import re
            raw_names = entry.get('distributors_raw', [])
            clean_names = [re.sub(r'\s*\(.*?\)', '', name) for name in raw_names]
            dists_text = "│ ".join(clean_names) + "│"
            
            dists = convert_to_bijoy(dists_text)
            set_cell_text(row.cells[2], dists)
            set_cell_text(row.cells[3], f"{entry['odo_start']}")
            set_cell_text(row.cells[4], f"{entry['odo_end']}")
            set_cell_text(row.cells[5], f"{entry['total_km']}")
            
            if entry.get('petrol_liters'):
                set_cell_text(row.cells[6], f"{entry['petrol_liters']}")
                set_cell_text(row.cells[7], f"{entry['petrol_cost']}/-")
            
            if entry.get('mobil_liters'):
                set_cell_text(row.cells[8], f"{entry['mobil_cost']}/-")
            
            set_cell_text(row.cells[9], f"{entry['da_amount']}/-")
            set_cell_text(row.cells[10], f"{entry['total_cost']}/-")
            
            if entry.get('others_designation'):
                set_cell_text(row.cells[11], convert_to_bijoy(entry['others_designation']))

    def _fill_total_row(self, row, summary):
        set_cell_text(row.cells[2], "‡gvU=")
        set_cell_text(row.cells[5], f"{summary['total_km']} wK:wg:")
        set_cell_text(row.cells[6], f"{summary['total_liters_petrol']} wjUvi")
        set_cell_text(row.cells[7], f"{summary['total_petrol_cost']}/-")
        set_cell_text(row.cells[8], f"{summary['total_mobil_cost']}/-")
        set_cell_text(row.cells[9], f"{summary['total_da']}/-")
        set_cell_text(row.cells[10], f"{summary['grand_total']}/")

    def _fill_summary(self, table, summary):
        # Left column (labels and values)
        set_cell_text(table.cell(0, 1), f"{summary['total_tour']} wU")
        set_cell_text(table.cell(1, 1), f"{summary['friday_tour']} wU" if summary['friday_tour'] > 0 else "bvB")
        set_cell_text(table.cell(2, 1), f"{summary['meeting_count']} wU")
        set_cell_text(table.cell(3, 1), f"{summary['manager_tour']} wU")
        set_cell_text(table.cell(4, 1), f"{summary['short_tour']} wU")
        set_cell_text(table.cell(5, 1), f"{summary['net_tours']} wU")

        # Right column (financials)
        set_cell_text(table.cell(0, 4), f"{summary['total_liters_petrol']} wjUvi")
        set_cell_text(table.cell(1, 4), f"{summary['total_km']} wKwg")
        set_cell_text(table.cell(2, 4), f"{summary['total_petrol_cost']} UvKv")
        set_cell_text(table.cell(3, 4), f"{summary['total_mobil_cost']} UvKv")
        set_cell_text(table.cell(4, 4), f"{summary['total_da']} UvKv")
        set_cell_text(table.cell(5, 4), f"{summary['total_others']} UvKv")
        set_cell_text(table.cell(6, 4), f"{summary['grand_total']} UvKv")
